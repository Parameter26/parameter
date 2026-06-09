from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/alwa/Documents/Parameter")
ASSETS = ROOT / "assets"
OUT = ROOT / "deliverables" / "instagram-content-plan"
DOCX_PATH = OUT / "Parameter_10_Konten_Awal_Instagram.docx"

NAVY = "0B1437"
NAVY_2 = "070B1F"
BLUE = "2563EB"
BLUE_BRIGHT = "3B82F6"
ICE = "EEF3FF"
SOFT = "F6F8FC"
BORDER = "E2E7F2"
GRAY = "55607F"
GRAY_LIGHT = "8A93AD"
WHITE = "FFFFFF"
GREEN = "16865C"
AMBER = "B56A00"


POSTS = [
    {
        "number": "01",
        "title": "Selamat Datang di Parameter",
        "format": "Carousel",
        "funnel": "Awareness",
        "role": "Pinned post",
        "objective": "Memperkenalkan masalah yang ingin diselesaikan Parameter dan membangun kedekatan dengan pemilik bisnis.",
        "hook": "Ada banyak bisnis bagus yang masih berjalan dengan cara yang melelahkan.",
        "slides": [
            "Ada banyak bisnis bagus yang masih berjalan dengan cara yang melelahkan.",
            "Data berantakan. Pekerjaan berulang. Pelanggan harus menunggu.",
            "Masalahnya bukan selalu kekurangan orang.",
            "Bisa jadi, sistemnya belum tepat.",
            "Parameter membantu bisnis bekerja lebih terarah lewat produk digital.",
            "Aplikasi, sistem, website, dan produk digital yang benar-benar digunakan.",
        ],
        "visual": "Gunakan komposisi editorial minimal. Slide 1 dominan tipografi; slide 2-4 memakai potongan UI, garis alur, atau daftar proses yang terasa semrawut; slide akhir menampilkan logo dan pixel mark.",
        "caption": (
            "Kami percaya teknologi yang baik seharusnya mengurangi pekerjaan yang melelahkan, "
            "bukan menambah kerumitan baru.\n\n"
            "Parameter membangun aplikasi, sistem internal, website, dan produk digital yang "
            "benar-benar digunakan. Kami juga menjalankan produk kami sendiri, jadi kami memahami "
            "bahwa pekerjaan tidak selesai ketika produk diluncurkan.\n\n"
            "Selamat datang di Parameter. Mari membangun sesuatu yang berguna."
        ),
        "cta": "Follow perjalanan Parameter dan lihat apa yang sedang kami bangun.",
    },
    {
        "number": "02",
        "title": "Apa Itu Parameter?",
        "format": "Carousel",
        "funnel": "Brand",
        "role": "Pinned post",
        "objective": "Menjelaskan posisi Parameter sebagai studio produk digital, bukan vendor proyek biasa.",
        "hook": "Kami bukan agency. Bukan juga software house biasa.",
        "slides": [
            "Kami bukan agency. Bukan juga software house biasa.",
            "Parameter adalah studio produk digital Indonesia.",
            "Kami membangun dan menjalankan produk kami sendiri.",
            "Pengalaman itu kami gunakan untuk membantu bisnis membangun sistemnya.",
            "Dari masalah, rancangan, produk, sampai benar-benar digunakan.",
            "Parameter — build products, improve systems.",
        ],
        "visual": "Pakai layout deklaratif dengan ruang putih lebar. Beri penekanan biru pada kata “studio produk digital”, “produk sendiri”, dan “digunakan”. Hindari ilustrasi generik.",
        "caption": (
            "Parameter adalah rumah bagi produk digital yang kami bangun sendiri dan layanan "
            "pengembangan untuk bisnis yang ingin bekerja lebih baik.\n\n"
            "Kami tidak hanya menerima brief lalu menyerahkan file. Kami ikut memahami masalah, "
            "merancang alur, membangun produk, dan mendampingi penggunaannya.\n\n"
            "Bagi kami, produk yang bagus bukan produk dengan fitur terbanyak. Produk yang bagus "
            "adalah produk yang dipakai."
        ),
        "cta": "Kenalan lebih jauh di parameter.cloud.",
    },
    {
        "number": "03",
        "title": "Yang Kami Bangun",
        "format": "Carousel",
        "funnel": "Consideration",
        "role": "Pinned post",
        "objective": "Memperlihatkan ekosistem produk dan layanan Parameter secara ringkas.",
        "hook": "Satu studio. Dua produk. Banyak masalah bisnis yang bisa dibereskan.",
        "slides": [
            "Satu studio. Dua produk. Banyak masalah bisnis yang bisa dibereskan.",
            "LMS Pro — sistem belajar untuk instruktur dan lembaga.",
            "Ritme — sistem operasional untuk studio kelas.",
            "Custom System — aplikasi dan dashboard sesuai proses bisnis.",
            "Website & Landing Page — mengubah perhatian menjadi tindakan.",
            "Satu parameter utama: harus benar-benar berguna.",
        ],
        "visual": "Tampilkan screenshot produk asli dalam frame sederhana. Setiap lini mendapat warna aksen sendiri, tetapi tetap memakai navy dan biru Parameter sebagai pengikat.",
        "caption": (
            "Kami membangun produk untuk menyelesaikan pekerjaan nyata.\n\n"
            "LMS Pro membantu instruktur dan lembaga mengelola pembelajaran. Ritme membantu studio "
            "kelas mengelola booking, member, credit, dan pembayaran. Untuk kebutuhan yang lebih "
            "spesifik, kami membangun sistem, aplikasi, website, dan landing page custom.\n\n"
            "Bentuknya boleh berbeda. Parameternya tetap sama: berguna, mudah dipakai, dan membantu "
            "bisnis bergerak."
        ),
        "cta": "Produk mana yang paling relevan dengan bisnismu? Tulis di komentar.",
    },
    {
        "number": "04",
        "title": "Bisnis Kamu Butuh Sistem?",
        "format": "Carousel",
        "funnel": "Education",
        "role": "Saveable",
        "objective": "Membantu audiens mengenali gejala operasional yang sudah membutuhkan sistem.",
        "hook": "5 tanda bisnis kamu sudah tidak cukup dikelola lewat chat dan spreadsheet.",
        "slides": [
            "5 tanda bisnis kamu sudah tidak cukup dikelola lewat chat dan spreadsheet.",
            "Data pelanggan tersebar di banyak tempat.",
            "Tim sering mengerjakan atau menanyakan hal yang sama.",
            "Pelanggan harus menunggu konfirmasi manual.",
            "Pemilik bisnis harus mengecek hampir semuanya.",
            "Laporan baru dibuat ketika sedang dibutuhkan.",
            "Kalau tiga atau lebih terasa familiar, prosesmu mungkin sudah butuh sistem.",
        ],
        "visual": "Gunakan pola checklist diagnostik. Masing-masing tanda divisualkan dengan contoh kecil: chat, sheet, reminder manual, approval, dan laporan.",
        "caption": (
            "Chat dan spreadsheet bukan alat yang buruk. Keduanya sangat berguna ketika bisnis masih "
            "sederhana.\n\n"
            "Masalah muncul ketika transaksi, pelanggan, dan anggota tim bertambah, tetapi prosesnya "
            "tetap bergantung pada ingatan dan pengecekan manual.\n\n"
            "Sistem dibutuhkan bukan agar bisnis terlihat canggih, melainkan agar pekerjaan penting "
            "tidak terus bergantung pada satu orang."
        ),
        "cta": "Simpan postingan ini, lalu hitung berapa tanda yang terjadi di bisnismu.",
    },
    {
        "number": "05",
        "title": "Website Bukan Sekadar Pajangan",
        "format": "Carousel / Reels",
        "funnel": "Education",
        "role": "Shareable",
        "objective": "Mengedukasi pasar bahwa website harus memiliki fungsi bisnis dan jalur konversi.",
        "hook": "Website yang cantik belum tentu bekerja.",
        "slides": [
            "Website yang cantik belum tentu bekerja.",
            "Pengunjung harus memahami bisnismu dalam beberapa detik.",
            "Mereka perlu alasan untuk percaya.",
            "Mereka harus tahu langkah berikutnya.",
            "Website harus cepat dan nyaman dibuka dari HP.",
            "Dan bisnis harus bisa ditemukan ketika orang mencarinya.",
            "Desain menarik perhatian. Struktur menghasilkan tindakan.",
        ],
        "visual": "Gunakan perbandingan before/after pada struktur halaman, bukan mengejek desain tertentu. Sorot headline, proof, CTA, mobile view, dan performa.",
        "caption": (
            "Website bukan brosur digital yang selesai setelah tampilannya disetujui.\n\n"
            "Ia harus menjawab pertanyaan calon pelanggan, membangun rasa percaya, dan mengarahkan "
            "mereka menuju tindakan yang jelas: menghubungi, mendaftar, membeli, atau meminta "
            "penawaran.\n\n"
            "Tampilan tetap penting. Tetapi tanpa pesan dan struktur yang tepat, perhatian tidak "
            "akan berubah menjadi tindakan."
        ),
        "cta": "Kirim postingan ini ke teman yang sedang menyiapkan website bisnis.",
    },
    {
        "number": "06",
        "title": "Kenapa Kami Membuat LMS Pro?",
        "format": "Story carousel",
        "funnel": "Product",
        "role": "Proof",
        "objective": "Menghubungkan masalah pengguna dengan alasan lahirnya LMS Pro.",
        "hook": "Mengelola kelas online seharusnya tidak membutuhkan tujuh aplikasi.",
        "slides": [
            "Mengelola kelas online seharusnya tidak membutuhkan tujuh aplikasi.",
            "Jadwal disimpan di spreadsheet.",
            "Materi dikirim melalui chat.",
            "Pembayaran diperiksa satu per satu.",
            "Rekaman kelas kembali tercecer.",
            "Karena itu, kami membangun LMS Pro.",
            "Kelas, jadwal, materi, pembayaran, dan rekaman dalam satu sistem.",
        ],
        "visual": "Mulai dengan visual banyak tab/aplikasi yang terpecah, lalu transisi menuju satu dashboard LMS Pro. Gunakan screenshot nyata sebagai klimaks.",
        "caption": (
            "LMS Pro lahir dari masalah yang sederhana tetapi berulang: pengelolaan kelas online "
            "tersebar di terlalu banyak tempat.\n\n"
            "Kami menyatukan jadwal, kelas live, materi, pembayaran, dan rekaman agar instruktur "
            "dapat fokus pada pengalaman belajar, bukan pekerjaan administratif.\n\n"
            "Ini juga menjadi cara kami membuktikan prinsip Parameter: mulai dari masalah nyata, "
            "bangun secukupnya, lalu terus perbaiki berdasarkan penggunaan."
        ),
        "cta": "Lihat LMS Pro melalui link produk di parameter.cloud.",
    },
    {
        "number": "07",
        "title": "Di Balik Ritme",
        "format": "Reels demo",
        "funnel": "Product",
        "role": "Product demo",
        "objective": "Menunjukkan alur penggunaan Ritme dan dampaknya bagi studio kelas.",
        "hook": "Studio kelas bukan cuma soal membuat jadwal.",
        "slides": [
            "Hook: studio kelas bukan cuma soal membuat jadwal.",
            "Member memilih kelas dan melakukan booking.",
            "Credit atau paket otomatis diperiksa.",
            "Reminder WhatsApp terkirim sebelum kelas.",
            "Pembayaran dan kehadiran tercatat.",
            "Pemilik studio melihat laporan tanpa rekap manual.",
            "Ritme — sistem operasional untuk studio kelas.",
        ],
        "visual": "Screen recording vertikal 20-30 detik. Pakai zoom halus pada tindakan utama dan subtitle besar. Jangan menampilkan semua fitur sekaligus.",
        "caption": (
            "Di balik satu kelas yang berjalan lancar, ada booking, kapasitas, paket member, credit, "
            "reminder, pembayaran, dan laporan yang harus dikelola.\n\n"
            "Ritme membantu studio yoga, pilates, dance, dan kelas kebugaran menangani alur tersebut "
            "dalam satu sistem lokal yang lebih rapi.\n\n"
            "Lebih sedikit pekerjaan administratif. Lebih banyak waktu untuk member dan kualitas kelas."
        ),
        "cta": "Punya studio kelas? Kirim DM dengan kata “RITME”.",
    },
    {
        "number": "08",
        "title": "Tidak Semua Bisnis Membutuhkan Aplikasi",
        "format": "Carousel",
        "funnel": "Authority",
        "role": "Opinion",
        "objective": "Membangun kepercayaan lewat sudut pandang yang jujur dan tidak memaksakan solusi.",
        "hook": "Jujur saja: bisnis kamu mungkin belum membutuhkan aplikasi.",
        "slides": [
            "Jujur saja: bisnis kamu mungkin belum membutuhkan aplikasi.",
            "Kalau hanya perlu dikenal, mulai dari website.",
            "Kalau menjalankan satu kampanye, gunakan landing page.",
            "Kalau pekerjaan internal terus berulang, pertimbangkan dashboard.",
            "Kalau pelanggan membutuhkan akses mandiri, aplikasi mulai masuk akal.",
            "Teknologi terbaik bukan yang paling rumit.",
            "Teknologi terbaik adalah yang paling tepat.",
        ],
        "visual": "Buat decision tree sederhana dari kebutuhan menuju bentuk solusi. Gaya konsultatif, bukan promosi. Gunakan ikon garis minimal dan banyak ruang kosong.",
        "caption": (
            "Membangun aplikasi hanya karena terdengar lebih canggih dapat menjadi investasi yang "
            "mahal dan tidak terpakai.\n\n"
            "Solusi harus mengikuti masalah. Kadang bisnis hanya membutuhkan halaman yang meyakinkan. "
            "Kadang proses internal perlu dirapikan. Aplikasi baru masuk akal ketika pengguna memang "
            "membutuhkan interaksi dan akses yang berulang.\n\n"
            "Kami lebih suka merekomendasikan solusi kecil yang tepat daripada solusi besar yang mubazir."
        ),
        "cta": "Simpan panduan ini sebelum memilih solusi digital.",
    },
    {
        "number": "09",
        "title": "Cara Parameter Membangun Produk",
        "format": "Carousel",
        "funnel": "Trust",
        "role": "Process",
        "objective": "Mengurangi keraguan calon klien dengan memperjelas proses kerja Parameter.",
        "hook": "Produk yang bagus tidak dimulai dari daftar fitur.",
        "slides": [
            "Produk yang bagus tidak dimulai dari daftar fitur.",
            "01 — Dengar: memahami masalah dan target sebenarnya.",
            "02 — Rancang: menyederhanakan alur pengguna.",
            "03 — Bangun: mengerjakan inti yang paling penting.",
            "04 — Uji: memastikan produk dapat digunakan dengan nyaman.",
            "05 — Dampingi: memperbaiki berdasarkan penggunaan nyata.",
            "Kami tidak mengejar banyak fitur. Kami mengejar produk yang dipakai.",
        ],
        "visual": "Visualkan proses sebagai garis progres modular. Sertakan artefak nyata seperti potongan brief, wireframe, task board, dan product screen.",
        "caption": (
            "Kami memulai proyek dengan memahami pekerjaan yang perlu diperbaiki, bukan dengan "
            "menentukan teknologi yang terlihat paling menarik.\n\n"
            "Setelah alurnya jelas, kami membangun bagian yang paling penting terlebih dahulu, "
            "menguji penggunaannya, lalu mengembangkan produk berdasarkan kebutuhan nyata.\n\n"
            "Proses ini menjaga produk tetap fokus dan membuat investasi digital lebih masuk akal."
        ),
        "cta": "Punya proses bisnis yang ingin dirapikan? Ceritakan lewat DM.",
    },
    {
        "number": "10",
        "title": "Konsultasi Tanpa Bahasa Teknis",
        "format": "Founder video / Single post",
        "funnel": "Conversion",
        "role": "Offer",
        "objective": "Mengubah audiens awal menjadi percakapan konsultasi dengan hambatan serendah mungkin.",
        "hook": "Punya masalah bisnis, tetapi belum tahu teknologi apa yang dibutuhkan?",
        "slides": [
            "Punya masalah bisnis, tetapi belum tahu teknologi apa yang dibutuhkan?",
            "Ceritakan proses yang paling menghabiskan waktu.",
            "Kami bantu memetakan masalahnya.",
            "Bisa jadi kamu butuh website, landing page, atau sistem internal.",
            "Bisa juga kamu belum membutuhkan semuanya.",
            "Konsultasi awal gratis. Tanpa bahasa teknis yang membingungkan.",
        ],
        "visual": "Founder berbicara langsung ke kamera di ruang kerja nyata. Sisipkan B-roll diskusi, wireframe, dan produk. Tutup dengan layar navy, logo, dan kata kunci DM.",
        "caption": (
            "Kamu tidak perlu datang dengan daftar fitur atau istilah teknis.\n\n"
            "Cukup ceritakan pekerjaan yang paling sering terlambat, berulang, sulit dipantau, atau "
            "membuat pelanggan menunggu. Kami akan membantu melihat akar masalah dan bentuk solusi "
            "digital yang paling masuk akal.\n\n"
            "Konsultasi awal gratis dan tidak mengharuskan kamu memulai proyek."
        ),
        "cta": "Kirim kata “PARAMETER” melalui DM untuk mulai ngobrol.",
    },
]


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=NAVY, bold=False, italic=False, font="Arial"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def format_paragraph(paragraph, before=0, after=6, line=1.25, align=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_text(doc, text, size=10.5, color=NAVY, bold=False, italic=False,
             before=0, after=6, line=1.25, align=None, keep=False):
    p = doc.add_paragraph()
    format_paragraph(p, before, after, line, align)
    p.paragraph_format.keep_together = keep
    p.paragraph_format.keep_with_next = keep
    set_run(p.add_run(text), size, color, bold, italic)
    return p


def add_label(doc, text):
    return add_text(
        doc, text.upper(), size=8.5, color=BLUE, bold=True,
        before=10, after=3, line=1.0, keep=True
    )


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    p.paragraph_format.keep_with_next = True
    return p


def add_pixel_mark(paragraph, size=9):
    for color in (BLUE, BLUE_BRIGHT, NAVY):
        run = paragraph.add_run("■")
        set_run(run, size=size, color=color, bold=True)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run(run, size=8.5, color=GRAY_LIGHT)


def add_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    add_pixel_mark(p, 6.5)
    r = p.add_run("  PARAMETER  /  INSTAGRAM CONTENT PLAN")
    set_run(r, size=8, color=GRAY_LIGHT, bold=True)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "5")
    bottom.set(qn("w:color"), BORDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    footer = section.footer
    footer.is_linked_to_previous = False
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_widths(table, [6500, 2860])
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"val": "nil"})
        set_cell_margins(cell, 0, 0, 0, 0)
    left = table.cell(0, 0).paragraphs[0]
    format_paragraph(left, after=0, line=1.0)
    set_run(left.add_run("parameter.cloud  •  @parameter.id2026"), size=8.5, color=GRAY_LIGHT)
    add_page_number(table.cell(0, 1).paragraphs[0])


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for idx, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, NAVY, 10, 5),
    ]:
        style = doc.styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_metadata_strip(doc, post):
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [3120, 3120, 3120])
    values = [
        ("FORMAT", post["format"]),
        ("FUNNEL", post["funnel"]),
        ("PERAN", post["role"]),
    ]
    for i, (label, value) in enumerate(values):
        cell = table.cell(0, i)
        set_cell_shading(cell, SOFT)
        set_cell_margins(cell, 110, 150, 110, 150)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            start={"val": "single", "sz": "6", "color": BORDER},
            end={"val": "single", "sz": "6", "color": BORDER},
        )
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        format_paragraph(p, after=1, line=1.0)
        set_run(p.add_run(label + "\n"), size=7.5, color=GRAY_LIGHT, bold=True)
        set_run(p.add_run(value), size=9.5, color=NAVY, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label, text, fill=ICE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    set_cell_border(
        cell,
        start={"val": "single", "sz": "24", "color": accent},
        top={"val": "nil"},
        bottom={"val": "nil"},
        end={"val": "nil"},
    )
    p = cell.paragraphs[0]
    format_paragraph(p, after=2, line=1.2)
    set_run(p.add_run(label.upper() + "\n"), size=8, color=accent, bold=True)
    text_color = WHITE if fill in (NAVY, NAVY_2) else NAVY
    set_run(p.add_run(text), size=11.5, color=text_color, bold=True)


def add_carousel_table(doc, slides):
    table = doc.add_table(rows=1, cols=2)
    widths = [900, 8460]
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, title in enumerate(("SLIDE", "COPY UTAMA")):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, 90, 140, 90, 140)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": NAVY},
            bottom={"val": "single", "sz": "6", "color": NAVY},
            start={"val": "single", "sz": "6", "color": NAVY},
            end={"val": "single", "sz": "6", "color": NAVY},
        )
        p = cell.paragraphs[0]
        format_paragraph(p, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
        set_run(p.add_run(title), size=8, color=WHITE, bold=True)

    for i, slide in enumerate(slides, start=1):
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell, 90, 140, 90, 140)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                start={"val": "single", "sz": "4", "color": BORDER},
                end={"val": "single", "sz": "4", "color": BORDER},
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cells[0], ICE)
        p_num = cells[0].paragraphs[0]
        format_paragraph(p_num, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_run(p_num.add_run(f"{i:02d}"), size=8.5, color=BLUE, bold=True)
        p_copy = cells[1].paragraphs[0]
        format_paragraph(p_copy, after=0, line=1.12)
        set_run(p_copy.add_run(slide), size=9.3, color=NAVY, bold=(i == 1))
    return table


def add_caption_box(doc, caption):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, SOFT)
    set_cell_margins(cell, 150, 180, 150, 180)
    set_cell_border(
        cell,
        top={"val": "single", "sz": "6", "color": BORDER},
        bottom={"val": "single", "sz": "6", "color": BORDER},
        start={"val": "single", "sz": "6", "color": BORDER},
        end={"val": "single", "sz": "6", "color": BORDER},
    )
    p = cell.paragraphs[0]
    format_paragraph(p, after=0, line=1.18)
    blocks = caption.split("\n\n")
    for idx, block in enumerate(blocks):
        if idx:
            p.add_run().add_break()
            p.add_run().add_break()
        set_run(p.add_run(block), size=9.2, color=GRAY)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.add_run().add_picture(str(ASSETS / "parameter-logo.png"), width=Inches(3.25))

    add_text(doc, "INSTAGRAM CONTENT PLAN", size=10, color=BLUE, bold=True, after=12, line=1.0)
    add_text(doc, "10 Konten Awal\nuntuk Membangun Brand Parameter", size=30, color=NAVY, bold=True, after=18, line=1.05)
    add_text(
        doc,
        "Rangkaian konten peluncuran yang membangun identitas, otoritas, bukti produk, dan percakapan penjualan.",
        size=14, color=GRAY, after=28, line=1.35
    )
    add_callout(
        doc,
        "Brand thesis",
        "Builder produk digital, bukan sekadar vendor.",
        fill=NAVY,
        accent=BLUE_BRIGHT,
    )
    add_text(doc, "Disiapkan untuk", size=8.5, color=GRAY_LIGHT, bold=True, before=28, after=3)
    add_text(doc, "@parameter.id2026  •  parameter.cloud", size=11, color=NAVY, bold=True, after=4)
    add_text(doc, "Juni 2026", size=9.5, color=GRAY_LIGHT, after=0)
    doc.add_page_break()


def add_strategy_page(doc):
    add_text(doc, "01  /  ARAH KOMUNIKASI", size=9, color=BLUE, bold=True, after=8)
    add_text(doc, "Bangun kepercayaan sebelum meminta penjualan.", size=24, color=NAVY, bold=True, after=12, line=1.08)
    add_text(
        doc,
        "Sepuluh konten pertama harus membuat audiens memahami siapa Parameter, apa yang dibangun, "
        "cara berpikir tim, dan alasan mereka layak memulai percakapan.",
        size=11.5, color=GRAY, after=14, line=1.35
    )

    add_heading(doc, "Narasi utama", 2)
    add_callout(
        doc,
        "Positioning",
        "Parameter membangun dan menjalankan produk sendiri, lalu membawa pengalaman itu untuk membantu bisnis membangun sistem yang benar-benar dipakai.",
    )

    add_heading(doc, "Komposisi 10 konten", 2)
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [3120, 3120, 3120])
    blocks = [
        ("30%", "IDENTITAS", "Siapa Parameter dan apa yang dibangun."),
        ("40%", "EDUKASI", "Membantu pasar memahami masalah digital."),
        ("30%", "BUKTI & OFFER", "Produk, proses, dan ajakan konsultasi."),
    ]
    for idx, (percent, label, desc) in enumerate(blocks):
        cell = table.cell(0, idx)
        set_cell_shading(cell, NAVY if idx == 0 else SOFT)
        set_cell_margins(cell, 160, 170, 160, 170)
        set_cell_border(
            cell,
            top={"val": "single", "sz": "6", "color": BORDER},
            bottom={"val": "single", "sz": "6", "color": BORDER},
            start={"val": "single", "sz": "6", "color": BORDER},
            end={"val": "single", "sz": "6", "color": BORDER},
        )
        p = cell.paragraphs[0]
        format_paragraph(p, after=0, line=1.1)
        fg = WHITE if idx == 0 else NAVY
        muted = "AEB8D8" if idx == 0 else GRAY
        set_run(p.add_run(percent + "\n"), size=20, color=BLUE_BRIGHT if idx == 0 else BLUE, bold=True)
        set_run(p.add_run(label + "\n"), size=8, color=fg, bold=True)
        set_run(p.add_run(desc), size=8.8, color=muted)

    add_heading(doc, "Aturan kreatif", 2)
    rules = [
        ("Satu ide per konten", "Jangan memadatkan semua layanan ke dalam satu unggahan."),
        ("Bukti lebih kuat dari klaim", "Gunakan screenshot produk, alur, dan artefak kerja asli."),
        ("Bahasa manusia", "Jelaskan dampak bisnis tanpa istilah teknis yang tidak diperlukan."),
        ("CTA sesuai funnel", "Awareness meminta follow; edukasi meminta save; offer meminta DM."),
    ]
    for title, desc in rules:
        p = doc.add_paragraph()
        format_paragraph(p, after=5, line=1.2)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        set_run(p.add_run("■  "), size=8, color=BLUE)
        set_run(p.add_run(title + " — "), size=10, color=NAVY, bold=True)
        set_run(p.add_run(desc), size=10, color=GRAY)

    add_heading(doc, "Pinned posts", 2)
    add_text(
        doc,
        "Pin konten 01, 02, dan 03. Ketiganya membentuk halaman depan profil: masalah yang dibela, "
        "identitas Parameter, lalu portofolio produk dan layanan.",
        size=10, color=GRAY, after=0
    )
    doc.add_page_break()


def add_post_page(doc, post, page_break_before=False):
    p = doc.add_paragraph()
    format_paragraph(p, after=4, line=1.0)
    p.paragraph_format.page_break_before = page_break_before
    set_run(p.add_run(post["number"] + "  /  CONTENT BLUEPRINT"), size=9, color=BLUE, bold=True)
    title = doc.add_paragraph()
    format_paragraph(title, after=10, line=1.03)
    title.paragraph_format.keep_with_next = True
    set_run(title.add_run(post["title"]), size=23, color=NAVY, bold=True)
    add_metadata_strip(doc, post)

    add_label(doc, "Tujuan")
    add_text(doc, post["objective"], size=10, color=GRAY, after=5, line=1.23)
    add_callout(doc, "Hook", post["hook"])

    add_label(doc, "Copy per slide / scene")
    add_carousel_table(doc, post["slides"])

    add_label(doc, "Arahan visual")
    add_text(doc, post["visual"], size=9.5, color=GRAY, after=4, line=1.2)

    add_label(doc, "Caption siap pakai")
    add_caption_box(doc, post["caption"])

    add_label(doc, "Call to action")
    add_callout(doc, "CTA", post["cta"], fill="EAF7F1", accent=GREEN)


def add_execution_page(doc):
    p = add_text(doc, "FINAL  /  EKSEKUSI", size=9, color=BLUE, bold=True, after=8)
    p.paragraph_format.page_break_before = True
    add_text(doc, "Urutan tayang yang disarankan", size=24, color=NAVY, bold=True, after=12, line=1.08)
    add_text(
        doc,
        "Publikasikan tiga konten fondasi lebih dahulu, lalu selingi edukasi, bukti produk, dan penawaran agar profil tidak terasa seperti katalog.",
        size=11, color=GRAY, after=14, line=1.3
    )

    schedule = [
        ("Minggu 1", "01 · 02 · 03", "Perkenalan brand dan pinned posts"),
        ("Minggu 2", "04 · 06 · 05", "Masalah operasional, bukti LMS Pro, edukasi website"),
        ("Minggu 3", "07 · 08", "Demo Ritme dan opini konsultatif"),
        ("Minggu 4", "09 · 10", "Proses kerja dan ajakan konsultasi"),
    ]
    table = doc.add_table(rows=1, cols=3)
    set_table_widths(table, [1800, 2100, 5460])
    for i, heading in enumerate(("PERIODE", "KONTEN", "FOKUS")):
        cell = table.cell(0, i)
        set_cell_shading(cell, NAVY)
        set_cell_margins(cell, 100, 140, 100, 140)
        p = cell.paragraphs[0]
        format_paragraph(p, after=0, line=1.0)
        set_run(p.add_run(heading), size=8, color=WHITE, bold=True)
    for period, content, focus in schedule:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell, 110, 140, 110, 140)
            set_cell_border(
                cell,
                top={"val": "single", "sz": "4", "color": BORDER},
                bottom={"val": "single", "sz": "4", "color": BORDER},
                start={"val": "single", "sz": "4", "color": BORDER},
                end={"val": "single", "sz": "4", "color": BORDER},
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for cell, text, bold, color in [
            (cells[0], period, True, NAVY),
            (cells[1], content, True, BLUE),
            (cells[2], focus, False, GRAY),
        ]:
            p = cell.paragraphs[0]
            format_paragraph(p, after=0, line=1.15)
            set_run(p.add_run(text), size=9.3, color=color, bold=bold)

    add_heading(doc, "Checklist sebelum publikasi", 2)
    checklist = [
        "Ukuran carousel 1080 × 1350 px; Reels 1080 × 1920 px.",
        "Headline dapat dibaca tanpa membuka caption.",
        "Maksimal satu pesan utama pada setiap slide.",
        "Screenshot produk menggunakan data demo, bukan data pengguna.",
        "Caption memakai pemisah paragraf agar mudah dibaca di ponsel.",
        "CTA hanya meminta satu tindakan.",
        "Tiga pinned post memiliki cover yang terlihat sebagai satu seri.",
        "Balas komentar dan DM maksimal dalam satu hari kerja.",
    ]
    for item in checklist:
        p = doc.add_paragraph()
        format_paragraph(p, after=5, line=1.2)
        set_run(p.add_run("□  "), size=11, color=BLUE, bold=True)
        set_run(p.add_run(item), size=10, color=GRAY)

    add_heading(doc, "Parameter keberhasilan 30 hari pertama", 2)
    add_callout(
        doc,
        "Fokus awal",
        "Profile visits, saves, shares, qualified DM, dan jumlah percakapan konsultasi — bukan follower semata.",
        fill=NAVY,
        accent=BLUE_BRIGHT,
    )
    add_text(
        doc,
        "Dokumen ini adalah fondasi. Setelah 10 konten tayang, gunakan performa aktual untuk menentukan "
        "seri lanjutan: edukasi UMKM, build in public, studi kasus, dan demo produk.",
        size=9.5, color=GRAY, before=10, after=0
    )


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.32)
    configure_styles(doc)
    add_header_footer(section)

    add_cover(doc)
    add_strategy_page(doc)
    for idx, post in enumerate(POSTS):
        add_post_page(doc, post, page_break_before=(idx > 0))
    add_execution_page(doc)

    props = doc.core_properties
    props.title = "Parameter — 10 Konten Awal Instagram"
    props.subject = "Instagram content plan and creative blueprint"
    props.author = "Parameter"
    props.keywords = "Parameter, Instagram, content plan, social media"
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
